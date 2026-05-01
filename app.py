import streamlit as st
import io
from datetime import datetime
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from fpdf import FPDF

# ==========================================
# 1. THE LEGAL DATA DICTIONARY
# ==========================================
# This acts as the logic engine's brain, containing the high-fidelity legal text.

BASE_POLICY = {
    "1. PURPOSE": "This Acceptable Use Policy (“Policy”) establishes the foundational standards for the acceptable, safe, and lawful use of {company_name} information systems, networks, devices, and data. This Policy is intended to protect the confidentiality, integrity, and availability of Company information, define acceptable and prohibited behaviors, and reduce legal, operational, and cybersecurity risk.",
    "2. SCOPE": "This Policy applies to all employees, independent contractors, vendors, and third-party affiliates (“Users”) utilizing {company_name}-owned or managed systems, devices, networks, and data. It extends to any personal devices utilized to access Company infrastructure under an approved Bring Your Own Device (BYOD) framework.",
    "3. USER RESPONSIBILITIES": "Users bear the primary responsibility for maintaining the security of {company_name} resources. Users must utilize Company resources exclusively for authorized business purposes, safeguard all credentials and sensitive organizational information, and comply with all applicable local, state, and federal regulations. Any suspected security incident, breach, or violation of this Policy must be immediately reported to {security_contact}.",
    "8. MONITORING AND PRIVACY": "To ensure compliance, security, and operational efficiency, {company_name} explicitly reserves the right to monitor, log, and audit all system usage, network traffic, and data stored on Company systems at any time, without prior notice. Users shall have no expectation of privacy regarding their use of Company IT resources, networks, or communications platforms.",
    "9. ENFORCEMENT & DISCIPLINARY ACTION": "Violations of this Policy represent a serious breach of organizational trust and security. Depending on the severity of the infraction, violations may result in immediate suspension or termination of network access, formal disciplinary action up to and including termination of employment or contract, and civil or criminal legal action where applicable.",
    "10. INCIDENT REPORTING": "Users are required to report any unauthorized access, loss of hardware, suspected phishing attempts, or known policy violations within 24 hours of discovery. All reports must be submitted directly to {security_contact}.",
    "11. ACKNOWLEDGMENT": "By accessing {company_name} systems, the User acknowledges that they have read, comprehended, and agreed to comply strictly with all requirements outlined in this Acceptable Use Policy."
}

TIERED_POLICY = {
    "4. CORE IT RULES": {
        "Flexible": "Users are permitted limited, reasonable personal use of Company networks and devices, provided such use does not interfere with daily business operations, consume excessive network bandwidth, or violate applicable laws. Users may download software required for their roles with informal managerial approval.",
        "Standard": "Company-provided IT resources are strictly intended for business purposes. While incidental personal use is tolerated, it must not impact productivity or introduce security risks. Users must lock devices when unattended. Only authorized, commercially vetted software on the approved software list may be installed.",
        "High-Security": "All IT resources, networks, and hardware are restricted exclusively to authorized business operations. Zero personal use is permitted. Users must operate within a Zero-Trust framework. Installation of unapproved software, circumventing endpoint management, or unauthorized external data transfers will result in immediate network revocation."
    },
    "5. GENERATIVE AI USAGE": {
        "Flexible": "Users are permitted to utilize public Generative AI tools (e.g., ChatGPT, Claude) for brainstorming, drafting non-sensitive content, and general research. Users must not input Highly Confidential data into public AI models. Outputs must be reasonably reviewed before business application.",
        "Standard": "Users may only utilize Generative AI platforms explicitly approved by {company_name} IT operations. Under no circumstances may protected, proprietary, or regulated client data be entered into an AI prompt. All AI-generated content must be validated for accuracy and cannot be represented as verified human-generated fact.",
        "High-Security": "The use of public or unsanctioned Generative AI tools is strictly prohibited. AI usage is limited entirely to Enterprise-approved, secure-enclave instances where data is not utilized to train external models. Strict data classification rules apply, and all AI queries are subject to continuous logging and auditing."
    },
    "6. REMOTE WORK / VPN": {
        "Flexible": "Users operating remotely are encouraged to utilize secure, private Wi-Fi networks. The use of a Virtual Private Network (VPN) is recommended when accessing sensitive company documents or internal servers.",
        "Standard": "Users accessing {company_name} infrastructure from remote locations must connect via the approved Company VPN. Multi-Factor Authentication (MFA) is strictly enforced for all remote access points. Public Wi-Fi usage is prohibited unless operating through the secure VPN tunnel.",
        "High-Security": "Remote access to {company_name} systems is heavily restricted and monitored. Users must utilize only Company-issued, managed hardware connecting via an always-on Enterprise VPN. Continuous endpoint detection and response (EDR) agents must remain active at all times during remote sessions."
    },
    "7. PERSONAL DEVICE / BYOD": {
        "Flexible": "Users may utilize personal devices (smartphones, tablets) to access Company email and communication platforms, provided the device maintains basic security controls, including a complex passcode and active biometric locks.",
        "Standard": "The use of personal devices to access {company_name} data requires formal device registration. Users must consent to the installation of Mobile Device Management (MDM) profiles to separate business and personal data. The Company reserves the right to remotely wipe business data from the personal device if lost or compromised.",
        "High-Security": "Bring Your Own Device (BYOD) is strictly prohibited. Company data, communications, and networks may only be accessed utilizing heavily managed, Company-issued hardware. Transferring Company data to personal devices, personal cloud storage, or personal email accounts is a severe security violation."
    }
}

# ==========================================
# 2. PDF TEXT SANITIZATION (Times core font / Latin-1)
# ==========================================
def sanitize_text(text: str) -> str:
    """Replace Unicode punctuation so FPDF core fonts accept the string."""
    s = text if isinstance(text, str) else str(text)
    return (
        s.replace("\u201c", '"')  # “
        .replace("\u201d", '"')  # ”
        .replace("\u2018", "'")  # ‘
        .replace("\u2019", "'")  # ’
        .replace("\u2014", "-")  # —
    )


# ==========================================
# 3. PDF GENERATOR CLASS (Legal Formatting)
# ==========================================
class LegalPDF(FPDF):
    def __init__(self, company_name):
        super().__init__()
        self.company_name = company_name

    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        # Arial italic 8
        self.set_font("Times", "I", 9)
        # The Custom Mojo Footer
        footer_text = f"Proprietary & Confidential | Prepared for {self.company_name} by SOP Mojo's AUP Engine."
        self.cell(0, 10, sanitize_text(footer_text), align="C")
# ==========================================
# 3. STREAMLIT UI & LOGIC
# ==========================================
st.set_page_config(page_title="SOP Mojo | AUP Engine", page_icon="⚡", layout="wide")

# Live token source (published Google Sheet)
SHEET_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vRrHS44BtJEmfCFcYEsfAs7V88mxrK5KVLVBSLxe-tUl84Y26DUjrHiusjundmCdAVDAYccTtdJSmpx/pubhtml"
SHEET_URL = SHEET_URL.strip().replace("/pubhtml", "/pub?output=csv")
orders_df = pd.read_csv(SHEET_URL, dtype=str)
VALID_TOKENS = (
    orders_df["Order ID"]
    .dropna()
    .str.replace(r'\.0$', '', regex=True)
    .str.strip()
    .tolist()
)
token_param = st.query_params.get("token")
provided_token = token_param[0] if isinstance(token_param, list) and token_param else token_param
provided_token = str(provided_token) if provided_token is not None else ""

if provided_token not in VALID_TOKENS:
    st.error("🛑 Access Denied: Invalid or missing authorization token.")
    st.markdown("[Purchase Access Here](https://samcart.com)")
    st.stop()

# Custom CSS to force the neon green text to be black for readability on buttons
st.markdown("""
    <style>
    div.stDownloadButton > button {
        color: #000000 !important;
        font-weight: 700 !important;
        border: none !important;
    }
    div.stDownloadButton > button:hover {
        opacity: 0.8;
    }
    div.stButton > button[kind="primary"] p {
        color: black !important;
        font-weight: 600 !important;
    }
    [data-testid="collapsedControl"] {
        color: #b0ff56 !important;
    }
    [data-testid="collapsedControl"] svg {
        fill: #b0ff56 !important;
        color: #b0ff56 !important;
        stroke: #b0ff56 !important;
    }
    </style>
""", unsafe_allow_html=True)

# Layout for Logo and Title
col_logo, col_title = st.columns([1, 8])
with col_logo:
    try:
        # Tries to load the logo if it exists in the folder
        st.image("logo.png", use_container_width=True)
    except:
        st.info("Logo Placeholder")

with col_title:
    st.title("Acceptable Use Policy (AUP) Engine")
    
st.markdown("Generate a legal-grade, custom Acceptable Use Policy for your organization in seconds. Select your compliance tiers and download a highly formatted document ready for signature.")

st.title("⚖️ Acceptable Use Policy (AUP) Engine")
st.markdown("Generate a legal-grade, custom Acceptable Use Policy for your organization in seconds. Select your compliance tiers and download a highly formatted document ready for signature.")

# Sidebar for Variables
with st.sidebar:
    st.header("🏢 Company Information")
    company_name = st.text_input("Legal Company Name", value="Acme Corp LLC")
    short_name = st.text_input("Short Name", value="")
    industry = st.text_input("Industry", value="Financial Services")
    policy_owner = st.text_input("Policy Owner (Name/Title)", value="Director of IT")
    security_contact = st.text_input("Security Contact Email", value="security@acmecorp.com")
    effective_date = st.date_input("Effective Date", value=datetime.today())

st.header("🛡️ Security & Compliance Tiers")
col1, col2 = st.columns(2)

with col1:
    core_tier = st.selectbox("1. Core IT Rules Tier", options=["Flexible", "Standard", "High-Security"], index=1)
    ai_tier = st.selectbox("2. Generative AI Usage Tier", options=["Flexible", "Standard", "High-Security"], index=1)

with col2:
    remote_tier = st.selectbox("3. Remote Work / VPN Tier", options=["Flexible", "Standard", "High-Security"], index=1)
    byod_tier = st.selectbox("4. Personal Device (BYOD) Tier", options=["Flexible", "Standard", "High-Security"], index=1)

def get_company_display_name():
    short = short_name.strip() if isinstance(short_name, str) else ""
    legal = company_name.strip() if isinstance(company_name, str) else ""
    return short or legal


def with_company(text, company_display):
    return text.replace("(Company)", company_display)


def build_policy_structure():
    return [
        {
            "type": "h1",
            "text": "Purpose",
            "children": [
                {
                    "type": "paragraph",
                    "text": "The purpose of the (Company) Acceptable Use Policy is to establish acceptable practices regarding the use of (Company) Information Resources in order to protect the confidentiality, integrity and availability of information created, collected, and maintained."
                }
            ],
        },
        {
            "type": "h1",
            "text": "Audience",
            "children": [
                {
                    "type": "paragraph",
                    "text": "The (Company) Acceptable Use Policy applies to any individual, entity, or process that interacts with any (Company) Information Resource."
                }
            ],
        },
        {
            "type": "h1",
            "text": "Policy",
            "children": [
                {
                    "type": "h2",
                    "text": "Acceptable Use",
                    "children": [
                        {"type": "paragraph", "text": "Personnel are responsible for complying with (Company) policies when using (Company) information resources and/or on (Company) time. If requirements or responsibilities are unclear, please seek assistance from the Information Security Committee."},
                        {"type": "paragraph", "text": "Personnel must promptly report harmful events or policy violations involving (Company) assets or information to their manager or a member of the Incident Handling Team. Events include, but are not limited to, the following:"},
                        {"type": "bullet", "text": "Technology incident: any potentially harmful event that may cause a failure, interruption, or loss in availability to (Company) Information Resources."},
                        {"type": "bullet", "text": "Data incident: any potential loss, theft, or compromise of (Company) information."},
                        {"type": "bullet", "text": "Unauthorized access incident: any potential unauthorized access to a (Company) Information Resource."},
                        {"type": "bullet", "text": "Facility security incident: any damage or potentially unauthorized access to a (Company) owned, leased, or managed facility."},
                        {"type": "bullet", "text": "Policy violation: any potential violation to this or other (Company) policies, standards, or procedures."},
                        {"type": "paragraph", "text": "Personnel should not purposely engage in activities that may harass, threaten, impersonate, or abuse others; degrade the performance of (Company) Information Resources; deprive authorized (Company) personnel access to a (Company) Information Resource; obtain additional resources beyond those allocated; or circumvent (Company) computer security measures."},
                        {"type": "paragraph", "text": "Personnel should not download, install, or run security programs or utilities that reveal or exploit weakness in the security of a system. For example, (Company) personnel should not run password cracking programs, packet sniffers, port scanners, or any other non-approved programs on any (Company) Information Resource."},
                        {"type": "paragraph", "text": "All inventions, intellectual property, and proprietary information, including reports, drawings, blueprints, software codes, computer programs, data, writings, and technical information, developed on (Company) time and/or using (Company) Information Resources are the property of (Company)."},
                        {"type": "paragraph", "text": "Use of encryption should be managed in a manner that allows designated (Company) personnel to promptly access all data."},
                        {"type": "paragraph", "text": "(Company) Information Resources are provided to facilitate company business and should not be used for personal financial gain."},
                        {"type": "paragraph", "text": "Personnel are expected to cooperate with incident investigations, including any federal or state investigations."},
                        {"type": "paragraph", "text": "Personnel are expected to respect and comply with all legal protections provided by patents, copyrights, trademarks, and intellectual property rights for any software and/or materials viewed, used, or obtained using (Company) Information Resources."},
                        {"type": "paragraph", "text": "Personnel should not intentionally access, create, store or transmit material which (Company) may deem to be offensive, indecent, or obscene."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Access Management",
                    "children": [
                        {"type": "bullet", "text": "Access to information is based on a 'need to know'."},
                        {"type": "bullet", "text": "Personnel are permitted to use only those network and host addresses issued to them by (Company) IT and should not attempt to access any data or programs contained on (Company) systems for which they do not have authorization or explicit consent."},
                        {"type": "bullet", "text": "All remote access connections made to internal (Company) networks and/or environments must be made through approved, and (Company)-provided, virtual private networks (VPNs)."},
                        {"type": "bullet", "text": "Personnel should not divulge any access information to anyone not specifically authorized to receive such information, including IT support personnel."},
                        {"type": "paragraph", "text": "Personnel must not share their personal authentication information, including:"},
                        {"type": "bullet", "text": "Account passwords"},
                        {"type": "bullet", "text": "Personal Identification Numbers (PINs)"},
                        {"type": "bullet", "text": "Security Tokens (i.e. Smartcard)"},
                        {"type": "bullet", "text": "Multi-factor authentication information"},
                        {"type": "bullet", "text": "Access cards and/or keys"},
                        {"type": "bullet", "text": "Digital certificates"},
                        {"type": "bullet", "text": "Similar information or devices used for identification and authentication purposes"},
                        {"type": "bullet", "text": "Access cards and/or keys that are no longer required must be returned to physical security personnel."},
                        {"type": "bullet", "text": "Lost or stolen access cards, security tokens, and/or keys must be reported to physical security personnel as soon as possible."},
                        {"type": "bullet", "text": "A service charge may be assessed for access cards, security tokens, and/or keys that are lost, stolen, or are not returned."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Authentication/Passwords",
                    "children": [
                        {"type": "paragraph", "text": "All personnel are required to maintain the confidentiality of personal authentication information."},
                        {"type": "paragraph", "text": "Any group/shared authentication information must be maintained solely among the authorized members of the group."},
                        {"type": "paragraph", "text": "All passwords, including initial and/or temporary passwords, must be constructed, and implemented according to the following (Company) rules:"},
                        {"type": "bullet", "text": "Must meet all requirements including minimum length, complexity, and reuse history."},
                        {"type": "bullet", "text": "Must not be easily tied back to the account owner by using things like username, social security number, nickname, relative's names, birth date, etc."},
                        {"type": "bullet", "text": "Must not be the same passwords used for non-business purposes."},
                        {"type": "bullet", "text": "Unique passwords should be used for each system, whenever possible."},
                        {"type": "paragraph", "text": "User account passwords must not be divulged to anyone. (Company) support personnel and/or contractors should never ask for user account passwords."},
                        {"type": "paragraph", "text": "If the security of a password is in doubt, the password should be changed immediately."},
                        {"type": "paragraph", "text": "Personnel should not circumvent password entry with application remembering, embedded scripts or hard coded passwords in client software."},
                        {"type": "paragraph", "text": "Security tokens (i.e. Smartcard) must be returned on demand or upon termination of the relationship with (Company), if issued."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Clear Desk/Clear Screen",
                    "children": [
                        {"type": "bullet", "text": "Personnel should log off from applications or network services when they are no longer needed."},
                        {"type": "bullet", "text": "Personnel should log off or lock their workstations and laptops when their workspace is unattended."},
                        {"type": "bullet", "text": "Confidential or internal information should be removed or placed in a locked drawer or file cabinet when the workstation is unattended and at the end of the workday if physical access to the workspace cannot be secured by other means."},
                        {"type": "bullet", "text": "Personal items, such as phones, wallets, and keys, should be removed or placed in a locked drawer or file cabinet when the workstation is unattended."},
                        {"type": "bullet", "text": "File cabinets containing confidential information should be locked when not in use or when unattended."},
                        {"type": "bullet", "text": "Physical and/or electronic keys used to access confidential information should not be left on an unattended desk or in an unattended workspace if the workspace itself is not physically secured."},
                        {"type": "bullet", "text": "Laptops should be either locked with a locking cable or locked away in a drawer or cabinet when the work area is unattended or at the end of the workday if the laptop is not encrypted."},
                        {"type": "bullet", "text": "Passwords must not be posted on or under a computer or in any other physically accessible location."},
                        {"type": "bullet", "text": "Copies of documents containing confidential information should be immediately removed from printers and fax machines."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Data Security",
                    "children": [
                        {"type": "bullet", "text": "Personnel should use approved encrypted communication methods whenever sending confidential information over public computer networks (Internet)."},
                        {"type": "bullet", "text": "Confidential information transmitted via USPS or other mail service must be secured in compliance with the Information Classification and Management Policy."},
                        {"type": "bullet", "text": "Only authorized cloud computing applications may be used for sharing, storing, and transferring confidential or internal information."},
                        {"type": "bullet", "text": "Information must be appropriately shared, handled, transferred, saved, and destroyed, based on the information sensitivity."},
                        {"type": "bullet", "text": "Personnel should not have confidential conversations in public places or over insecure communication channels, open offices, and meeting places."},
                        {"type": "bullet", "text": "Confidential information must be transported either by an (Company) employee or a courier approved by IT Management."},
                        {"type": "bullet", "text": "All electronic media containing confidential information must be securely disposed. Please contact IT for guidance or assistance."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Email and Electronic Communication",
                    "children": [
                        {"type": "bullet", "text": "Auto-forwarding electronic messages outside the (Company) internal systems is prohibited."},
                        {"type": "bullet", "text": "Electronic communications should not misrepresent the originator or (Company)."},
                        {"type": "bullet", "text": "Personnel are responsible for the accounts assigned to them and for the actions taken with their accounts."},
                        {"type": "bullet", "text": "Accounts must not be shared without prior authorization from (Company) IT, with the exception of calendars and related calendaring functions."},
                        {"type": "bullet", "text": "Employees should not use personal email accounts to send or receive (Company) confidential information."},
                        {"type": "paragraph", "text": "Any personal use of (Company) provided email should not:"},
                        {"type": "bullet", "text": "Involve solicitation"},
                        {"type": "bullet", "text": "Be associated with any political entity, excluding the (Company) sponsored PAC"},
                        {"type": "bullet", "text": "Have the potential to harm the reputation of (Company)"},
                        {"type": "bullet", "text": "Forward chain emails"},
                        {"type": "bullet", "text": "Contain or promote anti-social or unethical behavior"},
                        {"type": "bullet", "text": "Violate local, state, federal, or international laws or regulations"},
                        {"type": "bullet", "text": "Result in unauthorized disclosure of (Company) confidential information"},
                        {"type": "bullet", "text": "Or otherwise violate any other (Company) policies"},
                        {"type": "bullet", "text": "Personnel should only send confidential information using approved secure electronic messaging solutions."},
                        {"type": "bullet", "text": "Personnel should use caution when responding to, clicking on links within, or opening attachments included in electronic communications."},
                        {"type": "bullet", "text": "Personnel should use discretion in disclosing confidential or internal information in Out of Office or other automated responses, such as employment data, internal telephone numbers, location information or other sensitive data."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Hardware and Software",
                    "children": [
                        {"type": "bullet", "text": "All hardware must be formally approved by IT Management before being connected to (Company) networks."},
                        {"type": "bullet", "text": "Software installed on (Company) equipment must be approved by IT Management and installed by (Company) IT personnel."},
                        {"type": "bullet", "text": "All (Company) assets taken off-site should be physically secured at all times."},
                        {"type": "bullet", "text": "Personnel traveling to a High-Risk location, as defined by FBI and Office of Foreign Asset control, must contact IT for approval to travel with corporate assets."},
                        {"type": "bullet", "text": "Employees should not allow family members or other non-employees to access (Company) Information Resources."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Internet",
                    "children": [
                        {"type": "paragraph", "text": "The Internet must not be used to communicate (Company) confidential or internal information, unless the confidentiality and integrity of the information is ensured and the identity of the recipient(s) is established."},
                        {"type": "paragraph", "text": "Use of the Internet with (Company) networking or computing resources must only be used for business-related activities. Unapproved activities include, but are not limited to:"},
                        {"type": "bullet", "text": "Recreational games"},
                        {"type": "bullet", "text": "Streaming media"},
                        {"type": "bullet", "text": "Personal social media"},
                        {"type": "bullet", "text": "Accessing or distributing pornographic or sexually oriented materials"},
                        {"type": "bullet", "text": "Attempting or making unauthorized entry to any network or computer accessible from the Internet"},
                        {"type": "bullet", "text": "Or otherwise violate any other (Company) policies"},
                        {"type": "paragraph", "text": "Access to the Internet from outside the (Company) network using a (Company) owned computer must adhere to all of the same policies that apply to use from within (Company) facilities."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Mobile Devices and Bring Your Own Device (BYOD)",
                    "children": [
                        {"type": "bullet", "text": "The use of a personally owned mobile device to connect to the (Company) network is a privilege granted to employees only upon formal approval of IT Management."},
                        {"type": "bullet", "text": "All personally owned laptops and/or workstations must have approved virus and spyware detection/protection software along with personal firewall protection active."},
                        {"type": "bullet", "text": "Mobile devices that access (Company) email must have a PIN or other authentication mechanism enabled."},
                        {"type": "bullet", "text": "Confidential information should only be stored on devices that are encrypted in compliance with the (Company) Encryption Standard."},
                        {"type": "bullet", "text": "(Company) confidential information should not be stored on any personally owned mobile device."},
                        {"type": "bullet", "text": "Theft or loss of any mobile device that has been used to create, store, or access confidential or internal information must be reported to the (Company) Security Team immediately."},
                        {"type": "bullet", "text": "All mobile devices must maintain up-to-date versions of all software and applications."},
                        {"type": "bullet", "text": "All personnel are expected to use mobile devices in an ethical manner."},
                        {"type": "bullet", "text": "Jail-broken or rooted devices should not be used to connect to (Company) Information Resources."},
                        {"type": "bullet", "text": "(Company) IT Management may choose to execute 'remote wipe' capabilities for mobile devices without warning."},
                        {"type": "bullet", "text": "In the event that there is a suspected incident or breach associated with a mobile device, it may be necessary to remove the device from the personnel's possession as part of a formal investigation."},
                        {"type": "bullet", "text": "All mobile device usage in relation to (Company) Information Resources may be monitored, at the discretion of (Company) IT Management."},
                        {"type": "bullet", "text": "(Company) IT support for personally owned mobile devices is limited to assistance in complying with this policy. (Company) IT support may not assist in troubleshooting device usability issues."},
                        {"type": "bullet", "text": "Use of personally owned devices must be in compliance with all other (Company) policies."},
                        {"type": "bullet", "text": "(Company) reserves the right to revoke personally owned mobile device use privileges in the event that personnel do not abide by the requirements set forth in this policy."},
                        {"type": "bullet", "text": "Texting or emailing while driving is not permitted while on company time or using (Company) resources. Only hands-free talking while driving is permitted, while on company time or when using (Company) resources."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Physical Security",
                    "children": [
                        {"type": "bullet", "text": "Photographic, video, audio, or other recording equipment, such as cameras and cameras in mobile devices, is not allowed in secure areas."},
                        {"type": "bullet", "text": "Personnel must display photo ID access card at all times while in the building."},
                        {"type": "bullet", "text": "Personnel must badge in and out of access-controlled areas. Piggy-backing, tailgating, door propping and any other activity to circumvent door access controls are prohibited."},
                        {"type": "bullet", "text": "Visitors accessing card-controlled areas of facilities must be accompanied by authorized personnel at all times."},
                        {"type": "bullet", "text": "Eating or drinking are not allowed in data centers. Caution must be used when eating or drinking near workstations or information processing facilities."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Privacy",
                    "children": [
                        {"type": "paragraph", "text": "Information created, sent, received, or stored on (Company) Information Resources are not private and may be accessed by (Company) IT employees at any time, under the direction of (Company) executive management and/or Human Resources, without knowledge of the user or resource owner."},
                        {"type": "paragraph", "text": "(Company) may log, review, and otherwise utilize any information stored on or passing through its Information Resources."},
                        {"type": "paragraph", "text": "Systems Administrators, (Company) IT, and other authorized (Company) personnel may have privileges that extend beyond those granted to standard business personnel. Personnel with extended privileges should not access files and/or other information that is not specifically required to carry out an employment-related task."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Removable Media",
                    "children": [
                        {"type": "bullet", "text": "The use of removable media for storage of (Company) information must be supported by a reasonable business case."},
                        {"type": "bullet", "text": "All removable media use must be approved by (Company) IT prior to use."},
                        {"type": "bullet", "text": "Personally owned removable media use is not permitted for storage of (Company) information."},
                        {"type": "bullet", "text": "Personnel are not permitted to connect removable media from an unknown origin without prior approval from the (Company)."},
                        {"type": "bullet", "text": "Confidential and internal (Company) information should not be stored on removable media without the use of encryption."},
                        {"type": "bullet", "text": "All removable media must be stored in a safe and secure environment."},
                        {"type": "bullet", "text": "The loss or theft of a removable media device that may have contained any (Company) information must be reported to the (Company)."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Security Training and Awareness",
                    "children": [
                        {"type": "bullet", "text": "All new personnel must complete an approved security awareness training class prior to, or at least within 30 days of, being granted access to any (Company) Information Resources."},
                        {"type": "bullet", "text": "All personnel must be provided with and acknowledge they have received and agree to adhere to the (Company) Information Security Policies before they are granted to access to (Company) Information Resources."},
                        {"type": "bullet", "text": "All personnel must complete the annual security awareness training."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Social Media",
                    "children": [
                        {"type": "bullet", "text": "Communications made with respect to social media should be made in compliance with all applicable (Company) policies."},
                        {"type": "bullet", "text": "Personnel are personally responsible for the content they publish online."},
                        {"type": "bullet", "text": "Creating any public social media account intended to represent (Company), including accounts that could reasonably be assumed to be an official (Company) account, requires the permission of the (Company) Communications Departments."},
                        {"type": "paragraph", "text": "When discussing (Company) or (Company)-related matters, you should:"},
                        {"type": "bullet", "text": "Identify yourself by name"},
                        {"type": "bullet", "text": "Identify yourself as an (Company) representative"},
                        {"type": "bullet", "text": "Make it clear that you are speaking for yourself and not on behalf of (Company), unless you have been explicitly approved to do so"},
                        {"type": "bullet", "text": "Personnel should not misrepresent their role at (Company)."},
                        {"type": "paragraph", "text": "When publishing (Company)-relevant content online in a personal capacity, a disclaimer should accompany the content. An example disclaimer could be; 'The opinions and content are my own and do not necessarily represent (Company)'s position or opinion.'"},
                        {"type": "bullet", "text": "Content posted online should not violate any applicable laws (i.e. copyright, fair use, financial disclosure, or privacy laws)."},
                        {"type": "bullet", "text": "The use of discrimination in published content that is affiliated with (Company) will not be tolerated."},
                        {"type": "bullet", "text": "Confidential information, internal communications and non-public financial or operational information may not be published online in any form."},
                        {"type": "bullet", "text": "Personal information belonging to customers may not be published online."},
                        {"type": "bullet", "text": "Personnel approved to post, review, or approve content on (Company) social media sites must follow the (Company) Social Media Management Procedures."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "VoiceMail",
                    "children": [
                        {"type": "bullet", "text": "Personnel should use discretion in disclosing confidential or internal information in voicemail greetings, such as employment data, internal telephone numbers, location information or other sensitive data."},
                        {"type": "bullet", "text": "Personnel should not access another user's voicemail account unless it has been explicitly authorized."},
                        {"type": "bullet", "text": "Personnel must not disclose confidential information in voicemail messages."},
                    ],
                },
                {
                    "type": "h2",
                    "text": "Incidental Use",
                    "children": [
                        {"type": "paragraph", "text": "As a convenience to (Company) personnel, incidental use of Information Resources is permitted. The following restrictions apply:"},
                        {"type": "bullet", "text": "Incidental personal use of electronic communications, Internet access, fax machines, printers, copiers, and so on, is restricted to (Company) approved personnel; it does not extend to family members or other acquaintances."},
                        {"type": "bullet", "text": "Incidental use should not result in direct costs to (Company)."},
                        {"type": "bullet", "text": "Incidental use should not interfere with the normal performance of an employee's work duties."},
                        {"type": "bullet", "text": "No files or documents may be sent or received that may cause legal action against, or embarrassment to, (Company) or its customers."},
                        {"type": "bullet", "text": "Storage of personal email messages, voice messages, files and documents within (Company) Information Resources must be nominal."},
                        {"type": "bullet", "text": "All information located on (Company) Information Resources are owned by (Company) may be subject to open records requests and may be accessed in accordance with this policy."},
                    ],
                },
            ],
        },
        {
            "type": "h1",
            "text": "References",
            "children": [
                {"type": "bullet", "text": "ISO 27002: 6, 7, 8, 9, 11, 12, 13, 16, 18"},
                {"type": "bullet", "text": "NIST CSF: PR.AC, PR.AT, PR.DS, DE.CM, DE.DP, RS.CO"},
                {"type": "bullet", "text": "Asset Management Policy"},
                {"type": "bullet", "text": "Encryption Management Policy"},
                {"type": "bullet", "text": "Encryption Standard"},
                {"type": "bullet", "text": "Identity and Access Management Policy"},
                {"type": "bullet", "text": "Incident Management Policy"},
                {"type": "bullet", "text": "Information Classification and Management Policy"},
                {"type": "bullet", "text": "Mobile Device Acknowledgement"},
                {"type": "bullet", "text": "Personnel Security and Awareness Policy"},
                {"type": "bullet", "text": "Physical Security Policy"},
                {"type": "bullet", "text": "Social Media Management Procedure"},
            ],
        },
        {
            "type": "h1",
            "text": "Waivers",
            "children": [
                {"type": "paragraph", "text": "Waivers from certain policy provisions may be sought following the (Company) Waiver Process."}
            ],
        },
        {
            "type": "h1",
            "text": "Enforcement",
            "children": [
                {"type": "paragraph", "text": "Personnel found to have violated this policy may be subject to disciplinary action, up to and including termination of employment, and related civil or criminal penalties."},
                {"type": "paragraph", "text": "Any vendor, consultant, or contractor found to have violated this policy may be subject to sanctions up to and including removal of access rights, termination of contract(s), and related civil or criminal penalties."},
            ],
        },
    ]

# Document Generation Functions
def generate_docx():
    doc = Document()
    company_display = get_company_display_name()
    policy_structure = build_policy_structure()
    
    # --- TITLE PAGE ---
    # Add some spacing before title
    for _ in range(5):
        doc.add_paragraph()
        
    title = doc.add_paragraph("ACCEPTABLE USE POLICY")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    
    subtitle = doc.add_paragraph(f"{company_name.upper()}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    
    date_p = doc.add_paragraph(f"Effective Date: {effective_date.strftime('%B %d, %Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    
    # --- FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Proprietary & Confidential | Prepared for {company_name} by SOP Mojo's AUP Engine."
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.italic = True
    footer_para.runs[0].font.size = Pt(9)
    
    # --- CONTENT ---
    for section in policy_structure:
        heading = doc.add_heading(with_company(section["text"], company_display), level=1)
        heading.runs[0].font.size = Pt(14)
        for child in section["children"]:
            if child["type"] == "h2":
                subheading = doc.add_heading(with_company(child["text"], company_display), level=2)
                subheading.runs[0].font.size = Pt(12)
                for sub_item in child["children"]:
                    text = with_company(sub_item["text"], company_display)
                    if sub_item["type"] == "bullet":
                        para = doc.add_paragraph(text, style="List Bullet")
                    else:
                        para = doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                text = with_company(child["text"], company_display)
                if child["type"] == "bullet":
                    para = doc.add_paragraph(text, style="List Bullet")
                else:
                    para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf():
    pdf = LegalPDF(company_name=company_name)
    company_display = get_company_display_name()
    policy_structure = build_policy_structure()
    pdf.add_page()
    
    # --- TITLE PAGE ---
    pdf.ln(60)
    pdf.set_font("Times", "B", 24)
    pdf.cell(0, 15, sanitize_text("ACCEPTABLE USE POLICY"), align="C", ln=True)
    pdf.set_font("Times", "B", 16)
    pdf.cell(0, 10, sanitize_text(company_name.upper()), align="C", ln=True)
    pdf.ln(10)
    pdf.set_font("Times", "", 12)
    pdf.cell(
        0,
        10,
        sanitize_text(f"Effective Date: {effective_date.strftime('%B %d, %Y')}"),
        align="C",
        ln=True,
    )
    
    pdf.add_page()
    
    # --- CONTENT ---
    pdf.set_left_margin(10)
    policy_lines = []
    for section in policy_structure:
        policy_lines.append(f"## {with_company(section['text'], company_display)}")
        for child in section["children"]:
            if child["type"] == "h2":
                policy_lines.append(f"### {with_company(child['text'], company_display)}")
                for sub_item in child["children"]:
                    line_text = with_company(sub_item["text"], company_display)
                    if sub_item["type"] == "bullet":
                        policy_lines.append(f"* {line_text}")
                    else:
                        policy_lines.append(line_text)
            else:
                line_text = with_company(child["text"], company_display)
                if child["type"] == "bullet":
                    policy_lines.append(f"* {line_text}")
                else:
                    policy_lines.append(line_text)
        policy_lines.append("")

    for line in policy_lines:
        if line.startswith("## "):
            pdf.set_font("Times", "B", 14)
            pdf.set_x(10)
            pdf.multi_cell(0, 8, sanitize_text(line[3:].strip()))
            pdf.ln(1)
            continue

        if line.startswith("### "):
            pdf.set_font("Times", "B", 12)
            pdf.set_x(10)
            pdf.multi_cell(0, 7, sanitize_text(line[4:].strip()))
            pdf.ln(1)
            continue

        pdf.set_font("Times", "", 11)
        if line.startswith(" * "):
            pdf.set_x(25)
            text = line[3:].strip()
            pdf.multi_cell(0, 7, sanitize_text(text))
        elif line.startswith("* "):
            pdf.set_x(15)
            text = line[2:].strip()
            pdf.multi_cell(0, 7, sanitize_text(text))
        else:
            pdf.set_x(10)
            pdf.multi_cell(0, 7, sanitize_text(line))
        pdf.ln(1)
    
    # Output to byte array
    return bytes(pdf.output())

st.divider()

# Legal disclaimer (highly visible but unobtrusive)
st.caption(
    "Disclaimer: SOP Mojo is an operational frameworks provider, not a law firm. "
    "This tool generates a structural foundation for an Acceptable Use Policy based "
    "on standard industry practices. It does not constitute formal legal advice. We "
    "strongly recommend having your final document reviewed by qualified legal counsel "
    "prior to official company deployment."
)

if "documents_ready" not in st.session_state:
    st.session_state["documents_ready"] = False

if st.button("Generate Policy Documents", type="primary", use_container_width=True):
    required_business_details = [company_name, short_name, industry, policy_owner, security_contact]
    if any(not str(detail).strip() for detail in required_business_details):
        st.error("Action Required: Please fill out all business details before generating your policy.")
        st.session_state["documents_ready"] = False
        st.stop()
    st.session_state["documents_ready"] = True

if st.session_state["documents_ready"]:
    # Download Buttons
    col3, col4 = st.columns(2)

    with col3:
        docx_file = generate_docx()
        st.download_button(
            label="📝 Download Professional Word / Google Doc",
            data=docx_file,
            file_name=f"{company_name.replace(' ', '_')}_AUP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )

    with col4:
        pdf_file = generate_pdf()
        st.download_button(
            label="📥 Download Legal-Grade PDF",
            data=pdf_file,
            file_name=f"{company_name.replace(' ', '_')}_AUP.pdf",
            mime="application/pdf",
            use_container_width=True,
            type="primary"
        )